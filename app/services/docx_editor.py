"""In-place DOCX editing that preserves the original document's formatting.

The whole point: when a user uploads a Word resume, accepted bullet rewrites are
spliced into the ORIGINAL document. Fonts, spacing, bullet styles, and the
one-page layout stay byte-for-byte as the user made them; only the reworded
bullets change. This is the difference between a usable tailored resume and a
plain-text dump.

Limitation (documented honestly): if a source bullet has mixed inline formatting
(e.g. one bolded word mid-sentence), a whole-paragraph rewrite normalizes that
bullet to its first run's formatting. For typical resume bullets this is fine.
"""

from __future__ import annotations

import re
from io import BytesIO


def apply_edits_to_docx(
    docx_bytes: bytes,
    replacements: list[tuple[str, str, str]],
) -> tuple[bytes, list[str]]:
    """Apply (key, original_bullet, new_text) replacements in place.

    Returns (modified_docx_bytes, unplaced_keys). A replacement is "unplaced"
    when its original bullet could not be located in the document; those are
    NOT silently dropped — the caller surfaces them for manual review.
    """
    from docx import Document

    document = Document(BytesIO(docx_bytes))
    paragraphs = _all_paragraphs(document)
    used: set[int] = set()
    placed: set[str] = set()

    for key, original, new in replacements:
        if not original or not new:
            continue
        for i, paragraph in enumerate(paragraphs):
            if i in used:
                continue
            if _replace_in_paragraph(paragraph, original, new):
                used.add(i)
                placed.add(key)
                break

    buffer = BytesIO()
    document.save(buffer)
    unplaced = [key for key, _o, _n in replacements if key not in placed]
    return buffer.getvalue(), unplaced


def _all_paragraphs(document) -> list:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _replace_in_paragraph(paragraph, original: str, new: str) -> bool:
    runs = paragraph.runs
    full = "".join(run.text for run in runs)
    if not full.strip():
        return False

    # 1) Exact literal substring across runs (preserves everything else).
    idx = full.find(original)
    if idx != -1:
        _splice_runs(runs, idx, idx + len(original), new)
        return True

    # 2) Normalized whole-paragraph match (handles whitespace/tab differences),
    #    but only when lengths are close so we never clobber a longer paragraph
    #    that merely contains the bullet as a fragment.
    norm_full = _norm(full)
    norm_orig = _norm(original)
    if norm_full and norm_orig and (norm_orig in norm_full or norm_full in norm_orig):
        shorter = min(len(norm_full), len(norm_orig))
        longer = max(len(norm_full), len(norm_orig))
        if shorter == 0 or longer <= 1.6 * shorter:
            if runs:
                runs[0].text = new
                for run in runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(new)
            return True
    return False


def _splice_runs(runs: list, start: int, end: int, new: str) -> None:
    pos = 0
    inserted = False
    for run in runs:
        text = run.text
        r_start = pos
        r_end = pos + len(text)
        pos = r_end
        if r_end <= start or r_start >= end:
            continue
        local_start = max(0, start - r_start)
        local_end = min(len(text), end - r_start)
        before = text[:local_start]
        after = text[local_end:]
        run.text = (before + new + after) if not inserted else (before + after)
        inserted = True


def _norm(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()

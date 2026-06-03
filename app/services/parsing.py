from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO

from fastapi import UploadFile


MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
SUPPORTED_EXTENSIONS = {"pdf", "txt", "md", "docx"}


class UploadTooLarge(Exception):
    def __init__(self, size: int, limit: int) -> None:
        super().__init__(f"Uploaded CV is {size} bytes, limit is {limit}.")
        self.size = size
        self.limit = limit


class UnsupportedUpload(Exception):
    def __init__(self, suffix: str) -> None:
        super().__init__(f"Unsupported extension: '.{suffix}'")
        self.suffix = suffix


@dataclass(frozen=True)
class ParsedResume:
    text: str
    filename: str
    warnings: list[str]
    raw_bytes: bytes = b""
    ext: str = ""

    @property
    def is_docx(self) -> bool:
        return self.ext == "docx"


async def parse_resume_upload(upload: UploadFile | None) -> ParsedResume:
    if upload is None or not upload.filename:
        return ParsedResume(
            text="",
            filename="",
            warnings=["No CV/resume file was uploaded."],
        )

    content = await _read_capped(upload)
    filename = upload.filename
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if suffix in {"txt", "md"}:
        return ParsedResume(
            text=_decode_text(content),
            filename=filename,
            warnings=[],
            raw_bytes=content,
            ext=suffix,
        )

    if suffix == "pdf":
        parsed = _parse_pdf(content, filename)
        return ParsedResume(parsed.text, parsed.filename, parsed.warnings, raw_bytes=content, ext="pdf")

    if suffix == "docx":
        parsed = _parse_docx(content, filename)
        return ParsedResume(parsed.text, parsed.filename, parsed.warnings, raw_bytes=content, ext="docx")

    if suffix and suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedUpload(suffix)

    return ParsedResume(
        text=_decode_text(content),
        filename=filename,
        warnings=[
            f"No file extension detected. Treated '{filename}' as plain text.",
        ],
        raw_bytes=content,
        ext=suffix,
    )


async def _read_capped(upload: UploadFile) -> bytes:
    chunks: list[bytes] = []
    total = 0
    while True:
        chunk = await upload.read(64 * 1024)
        if not chunk:
            break
        total += len(chunk)
        if total > MAX_UPLOAD_BYTES:
            raise UploadTooLarge(total, MAX_UPLOAD_BYTES)
        chunks.append(chunk)
    return b"".join(chunks)


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace").strip()


def _parse_pdf(content: bytes, filename: str) -> ParsedResume:
    warnings: list[str] = []
    pypdf_text = _parse_pdf_with_pypdf(content, warnings)
    if pypdf_text:
        warnings.append("Review parsed CV text before trusting the analysis.")
        return ParsedResume(text=pypdf_text, filename=filename, warnings=warnings)

    warnings.append("pypdf returned no text; trying pdfplumber fallback.")
    pdfplumber_text = _parse_pdf_with_pdfplumber(content, warnings)
    if pdfplumber_text:
        warnings.append("Review parsed CV text before trusting the analysis.")
        return ParsedResume(text=pdfplumber_text, filename=filename, warnings=warnings)

    warnings.append(
        "PDF parsing returned no text after pypdf and pdfplumber. "
        "The CV is likely scanned/image-only; OCR or a text-based PDF is required."
    )
    return ParsedResume(text="", filename=filename, warnings=warnings)


def _parse_pdf_with_pypdf(content: bytes, warnings: list[str]) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        warnings.append("pypdf is not installed.")
        return ""

    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:  # noqa: BLE001 - surface parser failures clearly.
        warnings.append(f"pypdf parse failed: {exc}")
        return ""
    pages: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 - surface parser failures clearly.
            warnings.append(f"PDF page {index} parse failed: {exc}")
            page_text = ""
        pages.append(page_text)

    return "\n\n".join(page.strip() for page in pages if page.strip())


def _parse_docx(content: bytes, filename: str) -> ParsedResume:
    warnings: list[str] = []
    try:
        from docx import Document
    except ImportError:
        warnings.append("python-docx is not installed; treating .docx as opaque.")
        return ParsedResume(text="", filename=filename, warnings=warnings)

    try:
        document = Document(BytesIO(content))
    except Exception as exc:  # noqa: BLE001 - surface parser failures clearly.
        warnings.append(f"docx parse failed: {exc}")
        return ParsedResume(text="", filename=filename, warnings=warnings)

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    warnings.append("Review parsed CV text before trusting the analysis.")
    return ParsedResume(text="\n".join(parts), filename=filename, warnings=warnings)


def _parse_pdf_with_pdfplumber(content: bytes, warnings: list[str]) -> str:
    try:
        import pdfplumber
    except ImportError:
        warnings.append("pdfplumber is not installed.")
        return ""

    pages: list[str] = []
    try:
        with pdfplumber.open(BytesIO(content)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text(layout=True) or page.extract_text() or ""
                except Exception as exc:  # noqa: BLE001 - surface parser failures clearly.
                    warnings.append(f"pdfplumber page {index} parse failed: {exc}")
                    page_text = ""
                pages.append(page_text)
    except Exception as exc:  # noqa: BLE001 - surface parser failures clearly.
        warnings.append(f"pdfplumber parse failed: {exc}")
        return ""

    return "\n\n".join(page.strip() for page in pages if page.strip())

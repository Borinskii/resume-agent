"""Per-vacancy resume editor.

For each found vacancy the user can open a dedicated editor: vacancy-anchored
suggestions, a live text preview, and a download. When the source CV is a .docx,
the download is the ORIGINAL document with only the accepted bullets reworded
(formatting preserved). For PDF/text sources, the download is plain text because
PDF layout cannot be reconstructed honestly.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO

from app.services.encryption import EncryptionService
from app.services.storage import StoredAnalysis


@dataclass
class Suggestion:
    index: int
    skill_name: str
    role_title: str
    original_bullet: str
    suggested_bullet: str
    why: str
    provider: str
    accepted: bool
    edited_text: str

    @property
    def effective_text(self) -> str:
        return (self.edited_text or self.suggested_bullet).strip()


@dataclass
class GapItem:
    skill_name: str
    importance: str
    why_not: str


@dataclass
class RequirementRow:
    name: str
    importance: str
    evidence_level: int
    status: str
    quote: str


@dataclass
class EditorView:
    analysis_id: int
    role_index: int
    total_roles: int
    cv_filename: str
    source_ext: str
    format_preserved: bool
    format_note: str
    vacancy_title: str
    vacancy_company: str
    vacancy_provider: str
    vacancy_location: str
    vacancy_url: str
    current_score: int
    tailored_score: int
    requirements: list[RequirementRow] = field(default_factory=list)
    suggestions: list[Suggestion] = field(default_factory=list)
    gaps: list[GapItem] = field(default_factory=list)
    base_cv_text: str = ""
    edited_cv_text: str = ""
    unplaced_indices: list[int] = field(default_factory=list)
    has_llm_suggestions: bool = False
    rewrites_generated: bool = False


IMPORTANCE_LABEL = {
    "required": "a required skill",
    "preferred": "a preferred skill",
    "nice_to_have": "a nice-to-have",
}


def build_editor_view(
    record: StoredAnalysis,
    encryptor: EncryptionService,
    role_index: int,
    saved_edits: dict[int, dict],
    rewrites: list[dict] | None,
) -> EditorView:
    payload = record.payload
    roles = payload.get("roles") or []
    role_index = max(0, min(role_index, len(roles) - 1)) if roles else 0
    role = roles[role_index] if roles else {}

    try:
        base_cv_text = encryptor.decrypt(record.cv_ciphertext)
    except Exception:  # noqa: BLE001 - a broken key should not 500 the editor.
        base_cv_text = ""

    requirements = _requirement_rows(role)
    req_by_skill = {r.name.lower(): r for r in requirements}

    rewrites_generated = rewrites is not None
    suggestions: list[Suggestion] = []
    for index, rw in enumerate(rewrites or []):
        if rw.get("is_gap"):
            continue
        skill = rw.get("skill_name", "")
        saved = saved_edits.get(index, {})
        suggestions.append(
            Suggestion(
                index=index,
                skill_name=skill,
                role_title=rw.get("role_title", role.get("title", "")),
                original_bullet=rw.get("original_bullet", ""),
                suggested_bullet=rw.get("rewritten_bullet", ""),
                why=_why_for_skill(skill, role.get("title", ""), req_by_skill.get(skill.lower())),
                provider=rw.get("provider", ""),
                accepted=bool(saved.get("accepted")),
                edited_text=str(saved.get("edited_text") or ""),
            )
        )

    gaps = _gap_items(role)
    edited_cv_text, unplaced = apply_edits_text(base_cv_text, suggestions)

    source_ext = record.cv_ext or ""
    can_preserve = source_ext == "docx" and bool(record.cv_blob)
    format_note = _format_note(source_ext, can_preserve)

    return EditorView(
        analysis_id=record.id,
        role_index=role_index,
        total_roles=len(roles),
        cv_filename=record.cv_filename,
        source_ext=source_ext,
        format_preserved=can_preserve,
        format_note=format_note,
        vacancy_title=role.get("title", "Role"),
        vacancy_company=role.get("company", ""),
        vacancy_provider=role.get("provider", ""),
        vacancy_location=role.get("location", ""),
        vacancy_url=role.get("url", ""),
        current_score=int(role.get("current_score", 0)),
        tailored_score=int(role.get("tailored_score", 0)),
        requirements=requirements,
        suggestions=suggestions,
        gaps=gaps,
        base_cv_text=base_cv_text,
        edited_cv_text=edited_cv_text,
        unplaced_indices=unplaced,
        has_llm_suggestions=bool(suggestions),
        rewrites_generated=rewrites_generated,
    )


def _format_note(source_ext: str, can_preserve: bool) -> str:
    if can_preserve:
        return "Word source detected. The .docx download keeps your original layout, fonts, and one-page format; only accepted bullets change."
    if source_ext == "pdf":
        return "PDF source. Formatting cannot be reconstructed from PDF, so downloads are plain text. Upload the same resume as .docx to keep your layout."
    return "Plain-text source. Downloads are plain text. Upload a .docx to keep formatting."


def _requirement_rows(role: dict) -> list[RequirementRow]:
    rows: list[RequirementRow] = []
    for item in role.get("requirement_matches", []):
        rows.append(
            RequirementRow(
                name=item.get("name", ""),
                importance=item.get("importance", ""),
                evidence_level=int(item.get("evidence_level", 0)),
                status=item.get("status", ""),
                quote=item.get("quote", ""),
            )
        )
    rows.sort(key=lambda r: (r.importance != "required", r.evidence_level))
    return rows


def _why_for_skill(skill: str, role_title: str, req: RequirementRow | None) -> str:
    role = role_title or "this role"
    if req is None:
        return f"Surfaces {skill} more clearly for \"{role}\" by rephrasing an existing bullet. No new facts are added."
    importance = IMPORTANCE_LABEL.get(req.importance, req.importance)
    level_note = {
        0: "your CV has no clear evidence for it",
        1: "your CV only mentions it in passing",
        2: "your CV shows it in context",
        3: "your CV already demonstrates it well",
    }.get(req.evidence_level, "")
    return (
        f"\"{role}\" lists {skill} as {importance}. Right now {level_note} "
        f"(evidence level {req.evidence_level}). This rewrite surfaces {skill} for human "
        f"screeners and ATS keyword filters, without claiming anything not already in your CV."
    )


def _gap_items(role: dict) -> list[GapItem]:
    gaps: list[GapItem] = []
    seen: set[str] = set()
    by_name = {item.get("name", "").lower(): item for item in role.get("requirement_matches", [])}
    for name in role.get("missing_required", []):
        key = name.lower()
        if key in seen:
            continue
        seen.add(key)
        item = by_name.get(key, {})
        gaps.append(
            GapItem(
                skill_name=name,
                importance=item.get("importance", "required"),
                why_not=(
                    f"\"{role.get('title', 'The role')}\" requires {name}, but no supporting "
                    f"evidence was found in your CV. It is kept as a gap, not invented. "
                    f"Add it only if you genuinely have the experience."
                ),
            )
        )
    return gaps


def apply_edits_text(cv_text: str, suggestions: list[Suggestion]) -> tuple[str, list[int]]:
    """Plain-text splice for the live preview pane. Returns (text, unplaced indices)."""
    if not cv_text:
        accepted_now = [s for s in suggestions if s.accepted and s.effective_text]
        if not accepted_now:
            return "", []
        block = "\n".join(f"- {s.effective_text}" for s in accepted_now)
        return f"(No parsed CV text was available.)\n\nSuggested bullets:\n{block}", [s.index for s in accepted_now]

    lines = cv_text.split("\n")
    norm_lines = [_normalize_ws(line) for line in lines]
    used_line: set[int] = set()
    unplaced: list[int] = []

    for suggestion in suggestions:
        if not suggestion.accepted:
            continue
        target = suggestion.effective_text
        norm_orig = _normalize_ws(suggestion.original_bullet)
        if not target or not norm_orig:
            continue
        placed = False
        for i, norm_line in enumerate(norm_lines):
            if i in used_line or not norm_line:
                continue
            if norm_orig in norm_line or norm_line in norm_orig:
                lines[i] = target
                norm_lines[i] = _normalize_ws(target)
                used_line.add(i)
                placed = True
                break
        if not placed:
            unplaced.append(suggestion.index)

    edited = "\n".join(lines)
    if unplaced:
        by_index = {s.index: s for s in suggestions}
        review = "\n".join(f"- {by_index[i].effective_text}" for i in unplaced if i in by_index)
        edited += "\n\n--- Suggested edits that could not be auto-placed (review and insert manually) ---\n" + review
    return edited, unplaced


def accepted_replacements(view: EditorView) -> list[tuple[str, str, str]]:
    """(key, original_bullet, new_text) for accepted suggestions."""
    return [
        (str(s.index), s.original_bullet, s.effective_text)
        for s in view.suggestions
        if s.accepted and s.effective_text and s.original_bullet
    ]


def edited_resume_markdown(view: EditorView) -> str:
    lines = [f"# {view.cv_filename or 'Tailored CV'}"]
    if view.vacancy_title:
        anchor = view.vacancy_title
        if view.vacancy_company:
            anchor += f" at {view.vacancy_company}"
        lines.append(f"_Tailored for: {anchor}_")
    lines.append("")
    lines.append(view.edited_cv_text or "(No CV text available.)")
    lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def edited_resume_text(view: EditorView) -> str:
    return (view.edited_cv_text or "").rstrip() + "\n"


def edited_resume_docx(
    view: EditorView,
    record: StoredAnalysis,
    encryptor: EncryptionService,
) -> bytes:
    """Format-preserving when the source is .docx; otherwise a plain rebuild."""
    if view.format_preserved and record.cv_blob:
        from app.services.docx_editor import apply_edits_to_docx

        try:
            original = encryptor.decrypt_bytes(record.cv_blob)
        except Exception:  # noqa: BLE001
            original = b""
        if original:
            modified, _unplaced = apply_edits_to_docx(original, accepted_replacements(view))
            return modified

    return _plain_docx(view)


def _plain_docx(view: EditorView) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is required for DOCX export.") from exc

    document = Document()
    document.add_heading(view.cv_filename or "Tailored CV", level=1)
    if view.vacancy_title:
        anchor = view.vacancy_title
        if view.vacancy_company:
            anchor += f" at {view.vacancy_company}"
        document.add_paragraph(f"Tailored for: {anchor}")
    for chunk in (view.edited_cv_text or "").split("\n"):
        document.add_paragraph(chunk.rstrip())
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _normalize_ws(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()

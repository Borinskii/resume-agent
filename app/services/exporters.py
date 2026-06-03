from __future__ import annotations

from io import BytesIO

from app.services.encryption import EncryptionService
from app.services.storage import StoredAnalysis


def analysis_to_markdown(record: StoredAnalysis) -> str:
    payload = record.payload
    lines: list[str] = []
    lines.append(f"# CV gap report — {record.cv_filename or 'CV'}")
    lines.append(f"_Generated {record.created_at}_")
    lines.append("")
    lines.append("## Summary")
    lines.append(
        f"- Top role: **{payload.get('roles', [{}])[0].get('title', 'n/a') if payload.get('roles') else 'n/a'}**"
    )
    lines.append(f"- Average current fit: {payload.get('average_current', 0)}")
    lines.append(f"- Average tailored fit: {payload.get('average_tailored', 0)}")
    lines.append(f"- Confidence: {payload.get('confidence_score', 0)}")
    lines.append("")

    inventory = payload.get("skill_inventory") or []
    if inventory:
        lines.append("## Detected skills")
        lines.append(", ".join(inventory))
        lines.append("")

    for role in payload.get("roles", [])[:10]:
        lines.append(f"## {role.get('title', 'Role')}")
        if role.get("company"):
            lines.append(f"_{role.get('company')} · {role.get('provider', '')} · {role.get('location', '')}_")
        if role.get("url"):
            lines.append(f"[Open posting]({role['url']})")
        lines.append("")
        lines.append(
            f"- Current fit: {role.get('current_score', 0)} · Tailored: {role.get('tailored_score', 0)} · Confidence: {role.get('confidence_score', 0)}"
        )
        lines.append(
            f"- Required matched: {role.get('matched_required', 0)}/{role.get('total_required', 0)} · "
            f"Total: {role.get('matched_total', 0)}/{role.get('total_requirements', 0)}"
        )

        missing = role.get("missing_required") or []
        if missing:
            lines.append("")
            lines.append("**Required gaps**")
            for gap in missing:
                lines.append(f"- {gap}")

        actions = role.get("tailoring_actions") or []
        if actions:
            lines.append("")
            lines.append("**Tailoring actions**")
            for action in actions:
                lines.append(f"- _{action.get('kind', '').upper()}_: {action.get('title', '')}")
                if action.get("detail"):
                    lines.append(f"  - {action['detail']}")
                if action.get("source_quote"):
                    lines.append(f"  - > {action['source_quote']}")
        lines.append("")

    rewrites = record.rewrites
    if rewrites:
        lines.append("## Source-backed bullet rewrites")
        for item in rewrites:
            lines.append(f"### {item.get('skill_name', '')} — {item.get('role_title', '')}")
            lines.append(f"_Source bullet:_ {item.get('original_bullet', '')}")
            if item.get("is_gap"):
                lines.append(f"_Gap:_ {item.get('note', '')}")
            else:
                lines.append(f"_Rewrite:_ {item.get('rewritten_bullet', '')}")
            lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def tailored_cv_to_docx(record: StoredAnalysis, encryptor: EncryptionService) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is required for DOCX export.") from exc

    document = Document()
    document.add_heading(f"Tailored CV preview — {record.cv_filename or 'CV'}", level=1)
    document.add_paragraph(f"Generated {record.created_at}")

    document.add_heading("Original CV text", level=2)
    try:
        original = encryptor.decrypt(record.cv_ciphertext)
    except Exception:  # noqa: BLE001 - decryption failures fail the export.
        original = ""
    for chunk in original.split("\n"):
        if chunk.strip():
            document.add_paragraph(chunk)

    rewrites = record.rewrites or []
    if rewrites:
        document.add_heading("Recommended bullet rewrites (source-backed)", level=2)
        for item in rewrites:
            heading = f"{item.get('skill_name', '')} — {item.get('role_title', '')}"
            document.add_heading(heading, level=3)
            document.add_paragraph(f"Source bullet: {item.get('original_bullet', '')}")
            if item.get("is_gap"):
                document.add_paragraph(f"Gap: {item.get('note', '')}")
            else:
                document.add_paragraph(f"Rewrite: {item.get('rewritten_bullet', '')}")

    payload = record.payload
    missing = (payload.get("roles") or [{}])[0].get("missing_required") if payload.get("roles") else []
    if missing:
        document.add_heading("Required gaps for top role", level=2)
        for gap in missing:
            document.add_paragraph(gap, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

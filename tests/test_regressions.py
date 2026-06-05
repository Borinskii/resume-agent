"""Regression tests for gaps surfaced during the big testing pass:
DOCX tables, inline-bold preservation, revert, per-role isolation, segments, dedupe.
"""

from __future__ import annotations

import gc
import shutil
import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from app.services.docx_editor import apply_edits_to_docx
from app.services.editor import build_editor_view
from app.services.encryption import EncryptionService
from app.services.matching import analyze_resume
from app.services.storage import (
    DEFAULT_DB_PATH,
    StoredAnalysis,
    get_role_edits,
    get_role_rewrites,
    init_database,
    persist_analysis,
    save_role_edits,
    save_role_rewrites,
)


def _docx(build) -> bytes:
    from docx import Document
    doc = Document()
    build(doc)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _read(docx_bytes: bytes):
    from docx import Document
    return Document(BytesIO(docx_bytes))


class DocxTableTests(unittest.TestCase):
    def test_bullet_inside_table_is_replaced_and_table_survives(self):
        def build(doc):
            doc.add_paragraph("Header outside table.")
            t = doc.add_table(rows=1, cols=2)
            t.rows[0].cells[0].text = "Used Git and Docker in delivery workflows."
            t.rows[0].cells[1].text = "Other cell content."
        out, unplaced = apply_edits_to_docx(
            _docx(build),
            [("0", "Used Git and Docker in delivery workflows.", "Applied Git for version control.")],
        )
        d = _read(out)
        self.assertEqual(len(d.tables), 1)
        cells = [c.text for c in d.tables[0].rows[0].cells]
        self.assertIn("Applied Git for version control.", cells)
        self.assertIn("Other cell content.", cells)
        self.assertEqual(unplaced, [])


class DocxBoldPreservationTests(unittest.TestCase):
    def test_substring_replace_keeps_unrelated_bold_run(self):
        def build(doc):
            p = doc.add_paragraph()
            p.add_run("Skills: ")
            bold = p.add_run("Python")
            bold.bold = True
            p.add_run(", SQL, Java")
        out, unplaced = apply_edits_to_docx(_docx(build), [("0", ", SQL, Java", ", SQL, Java, C++")])
        d = _read(out)
        para = d.paragraphs[0]
        bold_text = "".join(r.text for r in para.runs if r.bold)
        self.assertIn("Python", bold_text)          # the bold run survived
        self.assertIn("C++", para.text)              # the edit landed
        self.assertEqual(unplaced, [])


class PerRoleIsolationTests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.mkdtemp()
        init_database(Path(self._tmp) / "app.sqlite")
        self._enc = EncryptionService(Path(self._tmp) / ".fernet_key")
        resume = "Built Python APIs with FastAPI and SQL queries."
        analysis = analyze_resume(resume_text=resume)
        self._aid = persist_analysis("u1", analysis, resume, self._enc, rewrites=[])

    def tearDown(self):
        gc.collect()
        shutil.rmtree(self._tmp, ignore_errors=True)
        init_database(DEFAULT_DB_PATH)

    def test_edits_are_isolated_per_role(self):
        save_role_edits(self._aid, 0, [{"index": 0, "accepted": True, "edited_text": "role0"}])
        save_role_edits(self._aid, 1, [{"index": 0, "accepted": False, "edited_text": "role1"}])
        e0 = get_role_edits(self._aid, 0)
        e1 = get_role_edits(self._aid, 1)
        self.assertTrue(e0[0]["accepted"])
        self.assertEqual(e0[0]["edited_text"], "role0")
        self.assertFalse(e1[0]["accepted"])
        self.assertEqual(e1[0]["edited_text"], "role1")

    def test_rewrites_are_isolated_per_role(self):
        save_role_rewrites(self._aid, 0, [{"skill_name": "Python"}])
        save_role_rewrites(self._aid, 1, [{"skill_name": "SQL"}, {"skill_name": "ML"}])
        self.assertEqual(len(get_role_rewrites(self._aid, 0)), 1)
        self.assertEqual(len(get_role_rewrites(self._aid, 1)), 2)
        self.assertIsNone(get_role_rewrites(self._aid, 2))


class EditorSegmentsAndRevertTests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.mkdtemp()
        init_database(Path(self._tmp) / "app.sqlite")
        self._enc = EncryptionService(Path(self._tmp) / ".fernet_key")

    def tearDown(self):
        gc.collect()
        shutil.rmtree(self._tmp, ignore_errors=True)
        init_database(DEFAULT_DB_PATH)

    def _record(self):
        cv = "Backend Engineer\nBuilt APIs with Python.\nUsed Docker in delivery."
        payload = {"roles": [{
            "title": "Backend Python Engineer", "company": "Co", "provider": "Greenhouse",
            "url": "https://x/1", "current_score": 50, "tailored_score": 60,
            "requirement_matches": [
                {"name": "Python", "importance": "required", "evidence_level": 1, "status": "mentioned", "quote": "Built APIs with Python."},
            ],
            "missing_required": [],
        }]}
        return StoredAnalysis(
            id=1, user_id="u1", created_at="2026-06-03T00:00:00+00:00", cv_sha256="x",
            cv_filename="cv.pdf", payload=payload, rewrites=[], cv_ciphertext=self._enc.encrypt(cv),
            top_score=50, top_title="Backend Python Engineer", cv_blob=b"", cv_ext="pdf",
        )

    def _rewrites(self):
        return [{"skill_name": "Python", "role_title": "Backend Python Engineer",
                 "original_bullet": "Built APIs with Python.",
                 "rewritten_bullet": "Engineered REST APIs in Python.",
                 "is_gap": False, "note": "", "provider": "ollama"}]

    def test_segments_have_anchor_for_matched_bullet(self):
        view = build_editor_view(self._record(), self._enc, 0, {}, self._rewrites())
        anchors = [s for s in view.segments if s["type"] == "anchor"]
        self.assertEqual(len(anchors), 1)
        self.assertEqual(anchors[0]["skill"], "Python")
        # Not accepted yet -> display shows the ORIGINAL line
        self.assertEqual(anchors[0]["display"], "Built APIs with Python.")

    def test_accepted_segment_shows_rewrite_then_revert_restores(self):
        accepted = build_editor_view(self._record(), self._enc, 0, {0: {"accepted": True, "edited_text": ""}}, self._rewrites())
        anchor = next(s for s in accepted.segments if s["type"] == "anchor")
        self.assertEqual(anchor["display"], "Engineered REST APIs in Python.")
        self.assertTrue(anchor["accepted"])
        # Revert == accepted False -> original restored
        reverted = build_editor_view(self._record(), self._enc, 0, {0: {"accepted": False, "edited_text": ""}}, self._rewrites())
        anchor2 = next(s for s in reverted.segments if s["type"] == "anchor")
        self.assertEqual(anchor2["display"], "Built APIs with Python.")
        self.assertFalse(anchor2["accepted"])


class TailoringDedupeTests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.mkdtemp()
        init_database(Path(self._tmp) / "app.sqlite")

    def tearDown(self):
        gc.collect()
        shutil.rmtree(self._tmp, ignore_errors=True)
        init_database(DEFAULT_DB_PATH)

    def test_same_bullet_yields_one_rewrite(self):
        from unittest import mock
        from app.services.tailoring import generate_rewrites_for_top_role

        # Two highlight actions whose quotes both map to the same CV line.
        top = {"title": "ML Engineer", "url": "",
               "tailoring_actions": [
                   {"kind": "highlight", "title": "Make Python more visible", "source_quote": "Programming Languages: Python, SQL"},
                   {"kind": "highlight", "title": "Make SQL more visible", "source_quote": "Programming Languages: Python, SQL"},
               ]}
        resume = "Skills\nProgramming Languages: Python, SQL, Java, C++\nExperience here."

        class _C:
            provider = "ollama"
            def complete(self, system, user):
                return "Proficient in Python and SQL across projects."

        with mock.patch("app.services.tailoring.build_llm_client", return_value=_C()):
            rewrites = generate_rewrites_for_top_role(top, resume)
        self.assertEqual(len(rewrites), 1)  # deduped to one card for that bullet


if __name__ == "__main__":
    unittest.main()

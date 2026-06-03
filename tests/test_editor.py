"""Unit tests for the per-vacancy resume editor and the DOCX format engine."""

from __future__ import annotations

import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from app.services.docx_editor import apply_edits_to_docx
from app.services.editor import (
    Suggestion,
    apply_edits_text,
    build_editor_view,
    edited_resume_docx,
    edited_resume_markdown,
)
from app.services.encryption import EncryptionService
from app.services.storage import DEFAULT_DB_PATH, StoredAnalysis, init_database


def _suggestion(index, original, suggested, accepted=True, edited_text=""):
    return Suggestion(
        index=index, skill_name="Python", role_title="Backend Engineer",
        original_bullet=original, suggested_bullet=suggested, why="because",
        provider="ollama", accepted=accepted, edited_text=edited_text,
    )


class ApplyEditsTextTests(unittest.TestCase):
    def test_exact_bullet_replaced(self):
        cv = "Backend Engineer\nBuilt APIs with Python.\nUsed Docker."
        edited, unplaced = apply_edits_text(cv, [_suggestion(0, "Built APIs with Python.", "Engineered REST APIs in Python.")])
        self.assertIn("Engineered REST APIs in Python.", edited)
        self.assertNotIn("Built APIs with Python.", edited)
        self.assertEqual(unplaced, [])

    def test_unaccepted_keeps_original(self):
        cv = "Built APIs with Python."
        edited, _ = apply_edits_text(cv, [_suggestion(0, "Built APIs with Python.", "X", accepted=False)])
        self.assertEqual(edited, cv)

    def test_edited_text_wins(self):
        cv = "Built APIs with Python."
        edited, _ = apply_edits_text(cv, [_suggestion(0, "Built APIs with Python.", "auto", edited_text="my Python wording.")])
        self.assertIn("my Python wording.", edited)
        self.assertNotIn("auto", edited)

    def test_unplaced_appended(self):
        cv = "Backend Engineer\nUsed Docker."
        edited, unplaced = apply_edits_text(cv, [_suggestion(0, "Nonexistent bullet.", "New text.")])
        self.assertIn("could not be auto-placed", edited)
        self.assertIn("New text.", edited)
        self.assertEqual(unplaced, [0])


class DocxInPlaceTests(unittest.TestCase):
    def _make_docx(self) -> bytes:
        from docx import Document
        doc = Document()
        doc.add_heading("Boris Resume", level=1)
        p = doc.add_paragraph()
        r1 = p.add_run("Built APIs with ")
        r2 = p.add_run("Python")
        r2.bold = True
        p.add_run(" and FastAPI.")
        doc.add_paragraph("Used Git and Docker in delivery workflows.")
        buf = BytesIO(); doc.save(buf); return buf.getvalue()

    def _text_of(self, docx_bytes: bytes) -> str:
        from docx import Document
        doc = Document(BytesIO(docx_bytes))
        return "\n".join(p.text for p in doc.paragraphs)

    def test_whole_paragraph_replaced_preserving_others(self):
        original = self._make_docx()
        new, unplaced = apply_edits_to_docx(
            original,
            [("0", "Used Git and Docker in delivery workflows.", "Applied Git for version control across delivery workflows.")],
        )
        text = self._text_of(new)
        self.assertIn("Applied Git for version control across delivery workflows.", text)
        self.assertNotIn("Used Git and Docker in delivery workflows.", text)
        # Heading and the other paragraph survive.
        self.assertIn("Boris Resume", text)
        self.assertIn("Built APIs with Python and FastAPI.", text)
        self.assertEqual(unplaced, [])

    def test_substring_replace_keeps_rest_of_paragraph(self):
        original = self._make_docx()
        new, unplaced = apply_edits_to_docx(original, [("0", "Built APIs with Python and FastAPI.", "Engineered Python and FastAPI services.")])
        text = self._text_of(new)
        self.assertIn("Engineered Python and FastAPI services.", text)
        self.assertEqual(unplaced, [])

    def test_unplaced_reported(self):
        original = self._make_docx()
        new, unplaced = apply_edits_to_docx(original, [("9", "This bullet is not in the document at all, nope.", "whatever")])
        self.assertEqual(unplaced, ["9"])

    def test_output_is_valid_docx_zip(self):
        original = self._make_docx()
        new, _ = apply_edits_to_docx(original, [("0", "Used Git and Docker in delivery workflows.", "Reworded.")])
        self.assertEqual(new[:2], b"PK")


class BuildEditorViewTests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.mkdtemp()
        init_database(Path(self._tmp) / "app.sqlite")
        self._enc = EncryptionService(Path(self._tmp) / ".fernet_key")

    def tearDown(self):
        import gc, shutil
        gc.collect()
        shutil.rmtree(self._tmp, ignore_errors=True)
        init_database(DEFAULT_DB_PATH)

    def _record(self, ext="pdf", blob=b""):
        cv = "Backend Engineer\nBuilt APIs with Python.\nUsed Docker in delivery."
        payload = {"roles": [
            {"title": "Backend Python Engineer", "company": "Example Co", "provider": "Greenhouse",
             "url": "https://example.com/1", "current_score": 50, "tailored_score": 60,
             "requirement_matches": [
                 {"name": "Python", "importance": "required", "evidence_level": 1, "status": "mentioned", "quote": "Built APIs with Python."},
                 {"name": "Kubernetes", "importance": "required", "evidence_level": 0, "status": "missing", "quote": ""},
             ],
             "missing_required": ["Kubernetes"]},
            {"title": "Data Analyst", "company": "Other", "provider": "Lever", "url": "", "current_score": 30, "tailored_score": 35,
             "requirement_matches": [], "missing_required": []},
        ]}
        return StoredAnalysis(
            id=1, user_id="u1", created_at="2026-06-03T00:00:00+00:00", cv_sha256="x",
            cv_filename="cv." + ext, payload=payload, rewrites=[], cv_ciphertext=self._enc.encrypt(cv),
            top_score=50, top_title="Backend Python Engineer", cv_blob=blob, cv_ext=ext,
        )

    def _rewrites(self):
        return [{"skill_name": "Python", "role_title": "Backend Python Engineer",
                 "original_bullet": "Built APIs with Python.",
                 "rewritten_bullet": "Engineered REST APIs in Python and FastAPI.",
                 "is_gap": False, "note": "", "provider": "ollama"}]

    def test_role_index_selects_vacancy(self):
        view = build_editor_view(self._record(), self._enc, 1, {}, [])
        self.assertEqual(view.role_index, 1)
        self.assertEqual(view.vacancy_title, "Data Analyst")
        self.assertEqual(view.total_roles, 2)

    def test_suggestion_has_why_tied_to_requirement(self):
        view = build_editor_view(self._record(), self._enc, 0, {}, self._rewrites())
        self.assertEqual(len(view.suggestions), 1)
        self.assertIn("Python", view.suggestions[0].why)
        self.assertIn("required", view.suggestions[0].why)

    def test_gap_listed(self):
        view = build_editor_view(self._record(), self._enc, 0, {}, self._rewrites())
        self.assertEqual([g.skill_name for g in view.gaps], ["Kubernetes"])

    def test_accepted_edit_applies_to_preview(self):
        saved = {0: {"accepted": True, "edited_text": ""}}
        view = build_editor_view(self._record(), self._enc, 0, saved, self._rewrites())
        self.assertIn("Engineered REST APIs in Python and FastAPI.", view.edited_cv_text)

    def test_pdf_source_not_format_preserved(self):
        view = build_editor_view(self._record(ext="pdf"), self._enc, 0, {}, self._rewrites())
        self.assertFalse(view.format_preserved)
        self.assertIn("plain text", view.format_note.lower())

    def test_docx_source_format_preserved(self):
        from docx import Document
        buf = BytesIO(); d = Document(); d.add_paragraph("Built APIs with Python."); d.save(buf)
        blob = self._enc.encrypt_bytes(buf.getvalue())
        rec = self._record(ext="docx", blob=blob)
        view = build_editor_view(rec, self._enc, 0, {0: {"accepted": True, "edited_text": ""}}, self._rewrites())
        self.assertTrue(view.format_preserved)
        out = edited_resume_docx(view, rec, self._enc)
        self.assertEqual(out[:2], b"PK")
        from docx import Document as D2
        text = "\n".join(p.text for p in D2(BytesIO(out)).paragraphs)
        self.assertIn("Engineered REST APIs in Python and FastAPI.", text)

    def test_markdown_export(self):
        view = build_editor_view(self._record(), self._enc, 0, {0: {"accepted": True, "edited_text": ""}}, self._rewrites())
        md = edited_resume_markdown(view)
        self.assertTrue(md.startswith("# "))
        self.assertIn("Engineered REST APIs", md)


if __name__ == "__main__":
    unittest.main()
